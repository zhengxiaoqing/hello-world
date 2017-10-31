import glob
import docx
#打开D盘picture下所有的图片，并保存为列表
listOfPicture = glob.glob(r'D:\picture\*.png')
#设置doc对象
doc = docx.Document()
#doc增加自然段
doc.add_paragraph('This is on the first page!', 'Title')
#循环列表
for p in listOfPicture:
	#插入标题
    doc.add_paragraph('test Picture' + str(listOfPicture.index(p)))
	#插入图片
    doc.add_picture(p, width = docx.shared.Cm(15), height = docx.shared.Cm(8))
#保存doc文档
doc.save('onePage.docx')
